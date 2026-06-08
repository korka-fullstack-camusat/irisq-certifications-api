"""
parser_service.py — Conversion universelle de documents → HTML + extraction de questions.

Formats supportés :
  • DOCX / DOC  → HTML fidèle (titres, gras/italique/souligné, listes ol/ul imbriquées,
                               tableaux, images base64, couleurs, alignement)
  • PDF         → HTML (texte extrait ligne par ligne, préservation de l'ordre)
  • TXT         → HTML (balise <pre> avec police monospace)
  • Autres      → chaîne vide (le frontend affichera un bouton de téléchargement)
"""

import base64
import os
import re
import uuid

import pdfplumber
from docx import Document


def _extract_pdf_page_text(page) -> str:
    """
    Extrait le texte d'une page PDF en filtrant les espaces fantômes de largeur nulle.

    Certaines polices PDF (ex: Cambria dans les modèles d'examen IRISQ) encodent les
    lettres accentuées (é, è, û, î…) comme la lettre nue suivie d'un caractère espace
    de largeur 0 — pdfplumber les restitue alors comme un vrai espace, ce qui coupe
    les mots en deux ("syste me" au lieu de "système"). On retire ces espaces fantômes
    avant extraction pour recoller les mots correctement.
    """
    filtered = page.filter(
        lambda obj: not (
            obj.get("object_type") == "char"
            and obj.get("text") == " "
            and obj.get("width", 1) == 0
        )
    )
    return filtered.extract_text() or ""


def _filtered_page_tables(page) -> list:
    """Tableaux d'une page PDF, en filtrant d'abord les espaces fantômes (cf. ci-dessus)."""
    filtered = page.filter(
        lambda obj: not (
            obj.get("object_type") == "char"
            and obj.get("text") == " "
            and obj.get("width", 1) == 0
        )
    )
    try:
        return filtered.extract_tables() or []
    except Exception:
        return []


def _extract_case_study_table_html(pdf) -> str | None:
    """
    Reconstruit fidèlement, en HTML, le tableau à compléter d'une étude de cas
    de type "Travail à faire" (ex : Section "Maîtrise des risques des SMQ" —
    colonnes Activité / Risque / Cause / Impact / P / G / C / Actions de maîtrise /
    Critère d'efficacité).

    Le texte linéarisé de pdfplumber casse complètement la mise en page de ce genre
    de tableau (cellules vides perdues, noms d'activités éclatés ligne par ligne…).
    On utilise donc l'extraction structurée `extract_tables()` — bien plus fiable
    pour ce cas — et on reconstitue : ligne d'en-têtes + une ligne par activité
    (nom pré-rempli, cellules vides à compléter par le candidat), exactement comme
    affiché dans le document source, y compris lorsque le tableau s'étale sur
    plusieurs pages.
    """
    HEADER_KEYWORDS = ["activit", "risque", "cause", "impact", "actions de ma", "crit"]
    headers: list[str] | None = None
    activities: list[str] = []
    header_page_idx: int | None = None

    def _join_wrapped(text: str) -> str:
        """
        Recolle un nom d'activité affiché sur plusieurs lignes dans une cellule étroite
        (ex: 'La\\nReception\\ndes\\nPatients' → 'La Reception des Patients'). Les retours
        à la ligne correspondent presque toujours à des mots séparés ; dans de rares cas
        un mot est coupé en milieu ('Consultatio'+'n des'), ce qui laisse une espace en
        trop — un compromis bien plus lisible qu'une concaténation hasardeuse erronée.
        """
        return re.sub(r"\s+", " ", text.replace("\n", " ")).strip()

    for page_idx, page in enumerate(pdf.pages):
        # Arrêter dès qu'on franchit l'en-tête de la section suivante
        if headers is not None and page_idx != header_page_idx:
            page_text = _extract_pdf_page_text(page)
            if re.search(r"^\s*Section\s+\d", page_text, re.IGNORECASE | re.MULTILINE):
                break

        for table in _filtered_page_tables(page):
            for row in table:
                cells = [(c or "").strip() for c in row]
                joined = " ".join(cells).lower()

                # Ligne d'en-têtes du tableau d'activités à risques
                if headers is None and all(k in joined for k in HEADER_KEYWORDS):
                    headers = []
                    for c in cells:
                        c = _fix_replacement_chars(_join_wrapped(c))
                        if c and c not in headers:
                            headers.append(c)
                    header_page_idx = page_idx
                    continue

                # Une fois l'en-tête trouvé : lignes "Activité" (1ère cellule renseignée,
                # le reste vide — colonnes à compléter par le candidat)
                if headers is not None:
                    first = cells[0] if cells else ""
                    rest_empty = all(not c for c in cells[1:])
                    if (
                        first and rest_empty
                        and not re.match(r"^(Page\s+\d|R[ée]f[ée]rence|Fiche\s+d|Implementor|Lead|Junior)", first, re.IGNORECASE)
                    ):
                        activities.append(_fix_replacement_chars(_join_wrapped(first)))

    if not headers or not activities:
        return None

    return _build_activity_table_html(headers, activities)


def _build_activity_table_html(headers: list[str], activities: list[str]) -> str:
    thead = "".join(
        f"<th style='border:1px solid #c7cbe0;padding:6px 8px;background:#e8eaf6;"
        f"font-size:0.78em;font-weight:700;color:#1a237e;text-align:left'>{_escape(h)}</th>"
        for h in headers
    )
    tbody = ""
    for activity in activities:
        cells_html = (
            f"<td style='border:1px solid #c7cbe0;padding:8px;font-weight:700;"
            f"vertical-align:top;background:#fafbff'>{_escape(activity)}</td>"
        )
        cells_html += "".join(
            "<td style='border:1px solid #c7cbe0;padding:8px;min-width:64px'>&nbsp;</td>"
            for _ in headers[1:]
        )
        tbody += f"<tr>{cells_html}</tr>"

    return (
        "<div style='overflow-x:auto;margin-top:10px'>"
        "<table style='border-collapse:collapse;width:100%;font-size:0.85em'>"
        f"<thead><tr>{thead}</tr></thead><tbody>{tbody}</tbody></table>"
        "</div>"
    )


def _extract_risk_legend_html(pdf) -> str | None:
    """
    Reconstruit fidèlement la légende "Probabilité (P) / Gravité (G)" + la matrice
    de criticité colorée (P × G), affichées juste avant le tableau d'activités dans
    la section "Maîtrise des risques des SMQ" — exactement comme dans le document
    source (cf. capture : table de définitions des niveaux 1/2/3, puis grille
    croisée Probabilité × Gravité avec les valeurs de criticité et leur code couleur
    vert/jaune/rouge).
    """
    levels: list[tuple[str, str, str]] = []
    found_header = False

    for page in pdf.pages:
        for table in _filtered_page_tables(page):
            for row in table:
                cells = [(c or "").strip().replace("\n", " ") for c in row]
                cells = [_fix_replacement_chars(re.sub(r"\s+", " ", c)) for c in cells]
                joined = " ".join(cells).lower().replace(" ", "")
                if not found_header and "probabilit" in joined and "gravit" in joined and "(p)" in joined:
                    found_header = True
                    continue
                if found_header and len(cells) >= 3 and re.match(r"^[123]$", cells[0]):
                    levels.append((cells[0], cells[1], cells[2]))
        if found_header and levels:
            break

    if not levels:
        return None

    rows_html = "".join(
        "<tr>"
        f"<td style='border:1px solid #c7cbe0;padding:6px 8px;font-weight:700;text-align:center;background:#e8eaf6'>{_escape(n)}</td>"
        f"<td style='border:1px solid #c7cbe0;padding:6px 8px'>{_escape(p)}</td>"
        f"<td style='border:1px solid #c7cbe0;padding:6px 8px'>{_escape(g)}</td>"
        "</tr>"
        for (n, p, g) in levels
    )
    legend_table = (
        "<table style='border-collapse:collapse;font-size:0.85em'>"
        "<thead><tr>"
        "<th style='border:1px solid #c7cbe0;padding:6px 8px;background:#e8eaf6'></th>"
        "<th style='border:1px solid #c7cbe0;padding:6px 8px;background:#e8eaf6;color:#1a237e;text-align:left'>Probabilité (P)</th>"
        "<th style='border:1px solid #c7cbe0;padding:6px 8px;background:#e8eaf6;color:#1a237e;text-align:left'>Gravité (G)</th>"
        "</tr></thead>"
        f"<tbody>{rows_html}</tbody>"
        "</table>"
    )

    # ── Matrice de criticité P × G (couleurs : vert ≤2, jaune 3-4, rouge ≥6,
    #     conformes au code couleur observé dans le document source) ──
    n = len(levels)

    def _crit_color(value: int) -> str:
        if value <= 2:
            return "#57bb8a"
        if value <= 4:
            return "#f7e379"
        return "#e8716d"

    header_cols = "".join(
        f"<th style='border:1px solid #c7cbe0;padding:4px 14px;background:#3f51b5;color:#fff;text-align:center'>{i}</th>"
        for i in range(1, n + 1)
    )
    matrix_rows = ""
    for g in range(1, n + 1):
        cells = "".join(
            f"<td style='border:1px solid #c7cbe0;padding:4px 14px;text-align:center;font-weight:700;"
            f"color:#1a237e;background:{_crit_color(p * g)}'>{p * g}</td>"
            for p in range(1, n + 1)
        )
        gravite_head = (
            f"<th rowspan='{n}' style='border:1px solid #c7cbe0;padding:4px 6px;background:#3f51b5;"
            f"color:#fff;text-align:center;white-space:nowrap'>Gravité</th>"
            if g == 1 else ""
        )
        matrix_rows += (
            "<tr>"
            f"{gravite_head}"
            f"<th style='border:1px solid #c7cbe0;padding:4px 14px;background:#3f51b5;color:#fff;text-align:center'>{g}</th>"
            f"{cells}"
            "</tr>"
        )

    matrix_table = (
        "<table style='border-collapse:collapse;font-size:0.85em'>"
        "<thead><tr>"
        "<th colspan='2' style='border:1px solid #c7cbe0;padding:4px;background:#fff'></th>"
        f"<th colspan='{n}' style='border:1px solid #c7cbe0;padding:4px 10px;background:#3f51b5;color:#fff;text-align:center'>Probabilité</th>"
        "</tr></thead>"
        f"<tbody>{matrix_rows}</tbody>"
        "</table>"
    )

    return (
        "<div style='margin:10px 0 16px'>"
        "<div style='display:flex;flex-wrap:wrap;gap:28px;align-items:flex-start'>"
        f"{legend_table}{matrix_table}"
        "</div>"
        "<p style='font-size:0.85em;margin-top:10px'><strong>P</strong> : Probabilité, "
        "<strong>G</strong> : Gravité, <strong>C</strong> : Criticité</p>"
        "</div>"
    )


def _extract_case_study_html(pdf) -> str | None:
    """
    Assemble le bloc HTML complet de l'étude de cas "Maîtrise des risques des SMQ" :
    légende Probabilité/Gravité + matrice de criticité, puis tableau d'activités à
    compléter — dans cet ordre, fidèle au document source.
    """
    legend_html   = _extract_risk_legend_html(pdf)
    activity_html = _extract_case_study_table_html(pdf)

    if not legend_html and not activity_html:
        return None

    return (legend_html or "") + (activity_html or "")


# ─────────────────────────────────────────────────────────────────────────────
# Utilitaires communs
# ─────────────────────────────────────────────────────────────────────────────

# Mots dont un caractère accentué est rendu "�" (remplacement Unicode irrécupérable)
# par l'encodage de la police Cambria-dérivée de ce modèle de document — on les
# corrige au cas par cas car ce sont toujours les mêmes mots qui apparaissent
# dans la légende et le tableau "Maîtrise des risques des SMQ".
_KNOWN_ACCENT_FIXUPS = {
    "Probabilit�": "Probabilité",
    "Gravit�": "Gravité",
    "Activit�": "Activité",
    "Crit�re": "Critère",
    "ma�trise": "maîtrise",
    "Ma�trise": "Maîtrise",
    "efficacit�": "efficacité",
    "Constat�": "Constaté",
    "L�analyse": "L'analyse",
    "L�": "L'",
}


def _fix_replacement_chars(text: str) -> str:
    """Corrige les mots connus dont un accent a été rendu '�' par pdfplumber."""
    if "�" not in text:
        return text
    for bad, good in _KNOWN_ACCENT_FIXUPS.items():
        text = text.replace(bad, good)
    return text


def _escape(text: str) -> str:
    """Échappe les caractères spéciaux HTML."""
    return (
        text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# ─────────────────────────────────────────────────────────────────────────────
# DOCX → HTML
# ─────────────────────────────────────────────────────────────────────────────

# Namespaces XML utilisés dans les fichiers OOXML
_W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def _build_image_cache(doc) -> dict:
    """
    Construit un dictionnaire rId → data-URI base64
    à partir des relations de la partie document principale.
    """
    cache: dict = {}
    try:
        for rId, rel in doc.part.rels.items():
            if "image" in (rel.reltype or "").lower():
                try:
                    part   = rel.target_part
                    b64    = base64.b64encode(part.blob).decode("utf-8")
                    ct     = part.content_type or "image/png"
                    cache[rId] = f"data:{ct};base64,{b64}"
                except Exception:
                    pass
    except Exception:
        pass
    return cache


def _extract_images(element, image_cache: dict) -> list:
    """
    Retourne la liste des balises <img> pour toutes les images
    référencées dans un élément XML (drawing / inline / anchor).
    """
    imgs = []
    # Recherche de tous les blip (références d'image DrawingML)
    blip_tag = f"{{{_A}}}blip"
    for blip in element.iter(blip_tag):
        rId = blip.get(f"{{{_R}}}embed")
        if rId and rId in image_cache:
            src = image_cache[rId]
            imgs.append(
                f'<img src="{src}" '
                f'style="max-width:100%;height:auto;display:block;margin:8px 0;'
                f'border-radius:4px" alt="image"/>'
            )
    return imgs


def _run_to_html(run, image_cache: dict) -> str:
    """Convertit un run DOCX en HTML (texte + images inline)."""
    html = ""

    # Vérifier s'il y a des images dans ce run
    imgs = _extract_images(run._r, image_cache)
    html += "".join(imgs)

    # Texte
    t = run.text
    if not t:
        return html

    t = _escape(t)

    # Couleur du texte (si définie et non noire)
    color = None
    try:
        if run.font and run.font.color and run.font.color.rgb:
            c = str(run.font.color.rgb).lower()
            if c not in ("000000", "auto"):
                color = f"#{c}"
    except Exception:
        pass

    # Mise en forme
    if run.bold:
        t = f"<strong>{t}</strong>"
    if run.italic:
        t = f"<em>{t}</em>"
    if run.underline:
        t = f"<u>{t}</u>"
    if color:
        t = f'<span style="color:{color}">{t}</span>'

    html += t
    return html


def _get_list_info(para) -> dict | None:
    """
    Retourne les infos de liste d'un paragraphe (niveau, type) ou None.
    """
    try:
        pPr = para._p.find(f"{{{_W}}}pPr")
        if pPr is None:
            return None
        numPr = pPr.find(f"{{{_W}}}numPr")
        if numPr is None:
            return None

        ilvl_el = numPr.find(f"{{{_W}}}ilvl")
        numId_el = numPr.find(f"{{{_W}}}numId")
        ilvl  = int(ilvl_el.get(f"{{{_W}}}val", 0)) if ilvl_el is not None else 0
        numId = int(numId_el.get(f"{{{_W}}}val", 0)) if numId_el is not None else 0

        if numId == 0:
            return None

        # Déterminer bullet vs numéroté via le style ou le numbering.xml
        style_name = (para.style.name or "").lower()
        is_bullet = "bullet" in style_name or "puce" in style_name

        # Essayer de lire le format depuis numbering.xml
        try:
            numbering_part = para.part.numbering_part
            if numbering_part is not None:
                num_el = numbering_part._element.find(
                    f".//{{{_W}}}num[@{{{_W}}}numId='{numId}']"
                )
                if num_el is not None:
                    abstractNumId_el = num_el.find(f"{{{_W}}}abstractNumId")
                    if abstractNumId_el is not None:
                        abId = abstractNumId_el.get(f"{{{_W}}}val")
                        abstract_el = numbering_part._element.find(
                            f".//{{{_W}}}abstractNum[@{{{_W}}}abstractNumId='{abId}']"
                        )
                        if abstract_el is not None:
                            lvl_el = abstract_el.find(
                                f".//{{{_W}}}lvl[@{{{_W}}}ilvl='{ilvl}']"
                            )
                            if lvl_el is not None:
                                numFmt_el = lvl_el.find(f"{{{_W}}}numFmt")
                                if numFmt_el is not None:
                                    fmt = numFmt_el.get(f"{{{_W}}}val", "")
                                    is_bullet = fmt in ("bullet", "none", "")
        except Exception:
            pass

        return {"level": ilvl, "numId": numId, "bullet": is_bullet}
    except Exception:
        return None


def _para_to_html(para, image_cache: dict) -> str:
    """Convertit un paragraphe DOCX complet en HTML."""
    style_name = (para.style.name or "") if para.style else ""

    # Extraire texte + images inline
    inner = "".join(_run_to_html(r, image_cache) for r in para.runs)

    # Images flottantes/block dans le paragraphe
    block_imgs = _extract_images(para._p, image_cache)
    imgs_html  = "".join(block_imgs)

    # Alignement
    align = ""
    try:
        fmt = para.paragraph_format
        if fmt.alignment is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if fmt.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                align = "text-align:center;"
            elif fmt.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                align = "text-align:right;"
            elif fmt.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                align = "text-align:justify;"
    except Exception:
        pass

    # Ligne vide
    if not inner.strip() and not imgs_html:
        return "<p style='margin:2px 0'>&nbsp;</p>"

    # Uniquement des images (pas de texte)
    if not inner.strip() and imgs_html:
        return imgs_html

    # ── Titres ──────────────────────────────────────────────────────────────
    base_style = (
        "font-weight:800;margin:1em 0 0.4em;color:#1a237e;"
        f"line-height:1.3;{align}"
    )
    if any(k in style_name for k in ("Heading 1", "Titre 1", "Title")):
        return f"<h2 style='font-size:1.25em;{base_style}'>{inner}</h2>{imgs_html}"
    if any(k in style_name for k in ("Heading 2", "Titre 2", "Subtitle")):
        return f"<h3 style='font-size:1.1em;{base_style}'>{inner}</h3>{imgs_html}"
    if any(k in style_name for k in ("Heading 3", "Titre 3")):
        return f"<h4 style='font-size:1em;{base_style}color:#283593'>{inner}</h4>{imgs_html}"
    if any(k in style_name for k in ("Heading 4", "Titre 4")):
        return f"<h5 style='font-size:0.95em;{base_style}color:#283593'>{inner}</h5>{imgs_html}"

    # ── Listes via numPr XML ─────────────────────────────────────────────────
    list_info = _get_list_info(para)
    if list_info:
        indent = list_info["level"] * 18
        bullet = "•" if list_info["bullet"] else ""
        marker = f'<span style="margin-right:6px">{bullet}</span>' if bullet else ""
        return (
            f"<div style='display:flex;align-items:baseline;margin:2px 0;"
            f"padding-left:{indent}px'>"
            f"{marker}"
            f"<span style='flex:1;{align}'>{inner}</span>"
            f"</div>{imgs_html}"
        )

    # ── Listes via nom de style ──────────────────────────────────────────────
    if any(k in style_name for k in ("List Bullet", "Puce")):
        return (
            f"<div style='display:flex;align-items:baseline;margin:2px 0;padding-left:16px'>"
            f"<span style='margin-right:6px'>•</span>"
            f"<span style='flex:1;{align}'>{inner}</span>"
            f"</div>{imgs_html}"
        )
    if any(k in style_name for k in ("List Number", "Numérotée")):
        return (
            f"<div style='display:flex;align-items:baseline;margin:2px 0;padding-left:16px'>"
            f"<span style='margin-right:6px'>◦</span>"
            f"<span style='flex:1;{align}'>{inner}</span>"
            f"</div>{imgs_html}"
        )

    # ── Paragraphe normal ────────────────────────────────────────────────────
    return (
        f"<p style='margin:0.35em 0;line-height:1.65;{align}'>{inner}</p>"
        f"{imgs_html}"
    )


def _table_to_html(table, image_cache: dict) -> str:
    """Convertit un tableau DOCX en tableau HTML avec fusion de cellules."""
    rows_html = []
    for row_idx, row in enumerate(table.rows):
        cells_html = []
        for cell in row.cells:
            # Contenu de la cellule (plusieurs paragraphes + sous-tableaux)
            cell_content = ""
            for child in cell._tc.iterchildren():
                local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if local == "p":
                    from docx.text.paragraph import Paragraph as P
                    cell_content += _para_to_html(P(child, cell), image_cache)
                elif local == "tbl":
                    from docx.table import Table as T
                    cell_content += _table_to_html(T(child, cell), image_cache)

            if row_idx == 0:
                style = (
                    "padding:7px 10px;border:1px solid #c5cae9;"
                    "font-weight:700;background:#e8eaf6;color:#1a237e;vertical-align:top"
                )
                tag = "th"
            else:
                style = (
                    "padding:6px 10px;border:1px solid #e0e0e0;"
                    "vertical-align:top;background:#fff"
                )
                tag = "td"

            cells_html.append(f"<{tag} style='{style}'>{cell_content}</{tag}>")

        rows_html.append(f"<tr>{''.join(cells_html)}</tr>")

    return (
        "<div style='overflow-x:auto;margin:0.8em 0'>"
        "<table style='border-collapse:collapse;width:100%;font-size:0.88em'>"
        + "".join(rows_html)
        + "</table></div>"
    )


def docx_to_html(file_path: str) -> str:
    """
    Convertit un fichier DOCX en HTML complet et fidèle.

    Supporte :
    - Titres Heading 1-4 (styles FR et EN)
    - Gras, italique, souligné, couleurs de texte
    - Listes à puces et numérotées (imbriquées)
    - Tableaux (y compris imbriqués)
    - Images inline et flottantes (encodées en base64)
    - Alignement (centre, droite, justifié)
    - Ordre original paragraphes + tableaux préservé
    """
    try:
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph

        doc         = Document(file_path)
        image_cache = _build_image_cache(doc)

        parts = [
            "<div style='font-family:inherit;font-size:0.9em;line-height:1.7;"
            "color:#1a1a2e;padding:4px 2px'>"
        ]

        for child in doc.element.body.iterchildren():
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if local == "p":
                para = DocxParagraph(child, doc)
                parts.append(_para_to_html(para, image_cache))

            elif local == "tbl":
                table = DocxTable(child, doc)
                parts.append(_table_to_html(table, image_cache))

            # Saut de page → ligne de séparation visuelle
            elif local == "sectPr":
                parts.append("<hr style='border:none;border-top:1px solid #e0e0e0;margin:12px 0'/>")

        parts.append("</div>")
        return "".join(parts)

    except Exception as e:
        print(f"[docx_to_html] Erreur : {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# PDF → HTML  (extraction texte + structure basique)
# ─────────────────────────────────────────────────────────────────────────────

def pdf_to_html(file_path: str) -> str:
    """
    Extrait le texte d'un PDF page par page et le retourne en HTML.
    Retourne "" si aucun texte n'est extractible (PDF scanné / image-only).
    Les PDFs sont alors affichables nativement via Blob URL dans un <iframe>.
    """
    try:
        has_text = False   # flag : au moins une ligne de texte réel trouvée
        parts = [
            "<div style='font-family:inherit;font-size:0.9em;line-height:1.7;"
            "color:#1a1a2e;padding:4px 2px'>"
        ]
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                if len(pdf.pages) > 1:
                    parts.append(
                        f"<p style='font-size:0.7em;font-weight:700;color:#9e9e9e;"
                        f"text-transform:uppercase;letter-spacing:0.1em;"
                        f"margin:12px 0 4px'>— Page {page_num} —</p>"
                    )
                text = _extract_pdf_page_text(page)
                for line in text.split("\n"):
                    esc = _escape(line)
                    if esc.strip():
                        parts.append(f"<p style='margin:0.25em 0'>{esc}</p>")
                        has_text = True   # texte réel détecté
                    else:
                        parts.append("<p style='margin:2px 0'>&nbsp;</p>")
        parts.append("</div>")

        # PDF scanné (image-only) : aucun texte extractible → retourner ""
        # Le frontend utilisera alors un Blob URL pour afficher le PDF en iframe.
        if not has_text:
            return ""

        return "".join(parts)
    except Exception as e:
        print(f"[pdf_to_html] Erreur : {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# TXT → HTML
# ─────────────────────────────────────────────────────────────────────────────

def txt_to_html(file_path: str) -> str:
    """Convertit un fichier texte brut en HTML."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = _escape(f.read())
        return (
            f"<div style='font-family:monospace;font-size:0.88em;line-height:1.6;"
            f"color:#1a1a2e;white-space:pre-wrap;padding:4px 2px'>{content}</div>"
        )
    except Exception as e:
        print(f"[txt_to_html] Erreur : {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# Convertisseur universel (point d'entrée principal)
# ─────────────────────────────────────────────────────────────────────────────

def document_to_html(file_path: str) -> str:
    """
    Convertit n'importe quel document supporté en HTML.

    Dispatche selon l'extension :
      .docx / .doc  → docx_to_html()   (rendu complet avec images)
      .pdf          → pdf_to_html()     (extraction texte page par page)
      .txt          → txt_to_html()     (pre-wrap monospace)
      autres        → ""               (le frontend affichera un lien de téléchargement)
    """
    if not os.path.exists(file_path):
        return ""

    ext = os.path.splitext(file_path)[1].lower()

    if ext in (".docx", ".doc"):
        return docx_to_html(file_path)
    elif ext == ".pdf":
        return pdf_to_html(file_path)
    elif ext in (".txt", ".text"):
        return txt_to_html(file_path)
    else:
        print(f"[document_to_html] Format non supporté : {ext}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# Parseur de questions interactif (mode Q&A structuré)
# ─────────────────────────────────────────────────────────────────────────────

def parse_exam_document(file_path: str) -> list:
    """
    Lit un PDF ou DOCX et tente d'extraire des questions/options
    selon un formatage numéroté (1. / A. / PARTIE …).
    Retourne [] si le document n'est pas dans ce format — dans ce cas
    exam_content_html prend le relais pour l'affichage.
    """
    if not os.path.exists(file_path):
        return []

    ext = os.path.splitext(file_path)[1].lower()
    raw_text = ""
    case_study_table_html = None

    try:
        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                raw_text = "\n".join(_extract_pdf_page_text(page) for page in pdf.pages)
                case_study_table_html = _extract_case_study_html(pdf)
        elif ext in (".doc", ".docx"):
            doc = Document(file_path)
            raw_text = "\n".join(p.text for p in doc.paragraphs)
        elif ext in (".txt", ".text"):
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                raw_text = f.read()
        else:
            return []
    except Exception as e:
        print(f"[parse_exam_document] Erreur lecture : {e}")
        return []

    return parse_exam_text(raw_text, case_study_table_html=case_study_table_html)


def _classify_line(line: str) -> str:
    """
    Classifie une ligne en catégorie : SKIP | SECTION | SUBSECTION |
    QUESTION | OPTION | JUSTIF | NOISE | TEXT
    """
    # ── NOISE : pointillés, lignes garbled (encoding PDF), numéros de page ──
    # Cas 1 : pointillés classiques
    if re.match(r"^[.\s…\-_]{5,}$", line):
        return "NOISE"
    # Cas 2 : caractères de remplacement Unicode � (dots encodés dans PDFs IRISQ)
    if len(line) >= 4 and sum(1 for c in line if c == "�") / len(line) > 0.5:
        return "NOISE"
    # Cas 3 : ligne ne contenant que des caractères non-ASCII répétés (garbled)
    stripped = line.replace(" ", "")
    if len(stripped) >= 5 and len(set(stripped)) <= 3 and not any(c.isalpha() for c in line):
        return "NOISE"
    if re.match(r"^(Page\s+\d|\d+\s*$)", line, re.IGNORECASE):
        return "NOISE"

    # ── SKIP : métadonnées de document (avec variantes d'encodage) ──
    if re.match(
        r"^(N.?\s*de\s*matricule|SUJET\s*:|Sujet\s+d.examen|Dur.e\s*:|"
        r"Documents?\s+[Aa]utoris|R.f.rence\s*:|Session\s+d.examen|"
        r"Fiche\s+d.examen|Concepteur|Approbateur|Page\s+\d+\s+sur|"
        r"PGE-|Module\s*:|Ressources\s+autoris|Niveau\s+de\s+comp|"
        r"Nom\s*:|Date\s*:|Signature)",
        line, re.IGNORECASE,
    ):
        return "SKIP"

    # ── SUMMARY : barème récapitulatif — ne pas parser comme questions ──
    if re.match(
        r"^(Bar[eè]me\s+r[eé]capitulatif|R[eé]capitulatif|"
        r"Total\s*:\s*\d+|TOTAL\s*:)",
        line, re.IGNORECASE,
    ):
        return "SUMMARY"

    # ── JUSTIF : demande de justification ──
    if re.match(
        r"^(Justifi[ez]|Motivez|Expliquez\s+votre|Commenter|"
        r"Justifiez\s+votre[/\s]?vos\s+choix)",
        line, re.IGNORECASE,
    ):
        return "JUSTIF"

    # ── SECTION : en-têtes de section majeure ──
    if re.match(
        r"^(Section\s+\d|SECTION\s+\d|Partie\s+[IVXivx\d]|PARTIE\s+[IVXivx\d]|"
        r"Chapitre\s+\d|CHAPITRE\s+\d|Module\s+\d|MODULE\s+\d|"
        r"[IVX]{1,4}\s*[.:\-]\s+.{3,}|"          # Chiffres romains : "I. Titre"
        r"EXERCICE\s+\d+\s*[:\-])",                # "EXERCICE 1 :"
        line, re.IGNORECASE,
    ):
        return "SECTION"

    # ── SUBSECTION : type de question ──
    if re.match(
        r"^(QCM|Q\.C\.M\.|Choix\s+[Mm]ultiple[s]?|"
        r"Questions?\s+[Oo]uvertes?|Questions?\s+[Rr][eé]dactionnelles?|"
        r"Questions?\s+[Àà]\s+d[eé]velopper|"
        r"Questions?\s+Vrai\s*[/\\]\s*Faux|"       # "Questions Vrai/Faux"
        r"Vrai\s*[/\\]\s*Faux|"                     # "Vrai/Faux" seul
        r"[EÉ]tude\s+de\s+[Cc]as|Cas\s+[Pp]ratique|"
        r"Travail\s+.{1,2}\s*[Ff]aire|APPLICATION|"
        r"Vrai\s+[Oo]u\s+[Ff]aux|True\s+or\s+False)",
        line, re.IGNORECASE,
    ):
        return "SUBSECTION"

    # ── OPTION : cases à cocher et lettres ──
    #  = caractère Wingdings checkbox utilisé dans les PDFs IRISQ (détecté par pdfplumber)
    # □☐○◦◻▪●• = variantes Unicode classiques
    # \uf07f = Wingdings checkbox in IRISQ PDFs, extracted as-is by pdfplumber
    if line and ord(line[0]) in (
        0xF07F, 0xF06F, 0xF0B7, 0xF0A7,  # Wingdings variants
        0x2022, 0x25A1, 0x2610, 0x25CF,  # Unicode bullets/boxes
        0x25E6, 0x25AA, 0x25CB, 0x25FB,
    ):
        return "OPTION"
    if re.match(r"^[□☐○◦◻▪●•]\s*.{1,}", line):
        return "OPTION"
    # Lettres A-D/a-d ou (A)-(D)
    if re.match(r"^[A-Ea-e][.)]\s+.{1,}", line):
        return "OPTION"
    if re.match(r"^\([A-Ea-e]\)\s+.{1,}", line):
        return "OPTION"
    # Tirets et puces introduisant une réponse courte (< 120 chars)
    if re.match(r"^[-–•]\s+.{2,}", line) and len(line) < 120:
        return "OPTION"

    # ── QUESTION : lignes numérotées ──
    if re.match(r"^(\d{1,2})[.)]\s+.{5,}", line):
        return "QUESTION"
    if re.match(r"^Q\s*(\d{1,2})\s*[.:)]\s*.{3,}", line, re.IGNORECASE):
        return "QUESTION"
    if re.match(r"^Question\s+\d+\s*[.:)]\s*.{3,}", line, re.IGNORECASE):
        return "QUESTION"
    if re.match(r"^(Cas|Exercice|Problem[e]?)\s+\d+\s*[.:)]\s*.+", line, re.IGNORECASE):
        return "QUESTION"

    return "TEXT"


def parse_exam_text(raw_text: str, case_study_table_html: str | None = None) -> list:
    """
    Parse universel : s'adapte automatiquement à tout template d'examen.

    Formats supportés (non exhaustif) :
      IRISQ Implementor  – Section X :, QCM : (pts), options □/☐
      IRISQ Lead         – Questions Vrai/Faux, sous-questions a/b/c + Vrai ☐ Faux ☐
      Standard           – PARTIE X, options A./B./C./D.
      Questions Ouvertes – questions rédactionnelles
      Libre              – numérotation 1./2./3.
    """
    questions          = []
    lines              = [l.strip() for l in raw_text.split("\n") if l.strip()]
    current_section    = ""
    current_subsection = ""
    current_q: dict | None = None
    # Pour les questions composées Vrai/Faux
    in_vraifaux_section = False
    current_part: dict | None = None   # sous-question courante (a/b/c)
    # Pour les études de cas : "Contexte N : <texte>" affiché avant chaque question liée
    current_context     = ""
    collecting_context  = False
    # Pour les études de cas basées sur un tableau à compléter ("Travail à Faire" + tableau,
    # ex: Section "Maîtrise des risques des SMQ") : tout le bloc (consignes + tableau)
    # est capturé comme une question ouverte unique, car il n'y a pas de questions numérotées
    collecting_travail  = False
    # Dès qu'on atteint la légende/tableau (zone reconstruite proprement en HTML via
    # table_html), on arrête de recopier le texte linéarisé brut (qui serait garbled)
    travail_table_started = False

    def _save():
        nonlocal current_q, current_part
        if current_q:
            q = current_q
            # Finaliser la dernière sous-question si composée
            if current_part and q.get("type") == "compound":
                if current_part not in q["parts"]:
                    q["parts"].append(current_part)
            current_part = None
            # Post-traitement QCM : si ≥ 2 options → confirmer type qcm
            if q.get("type") not in ("compound",) and len(q.get("options", [])) >= 2:
                q["type"] = "qcm"
            questions.append(q)
            current_q = None

    def _new_q(text: str, qtype: str = "open") -> dict:
        return {
            "id":                str(uuid.uuid4()),
            "section":           current_section,
            "subsection":        current_subsection,
            "part":              current_section or current_subsection or "",
            "type":              qtype,
            "text":              text,
            "options":           [],
            "parts":             [],      # sous-questions pour les composées
            "has_justification": False,
        }

    def _opt_text(line: str) -> str:
        """Extrait le texte propre d'une ligne d'option."""
        # Supprime les caractères Wingdings/checkbox en début (ord F07F, F06F…)
        clean = line
        while clean and ord(clean[0]) in (
            0xF07F, 0xF06F, 0xF0B7, 0xF0A7,
            0x2022, 0x25A1, 0x2610, 0x25CF,
            0x25E6, 0x25AA, 0x25CB, 0x25FB,
        ):
            clean = clean[1:]
        # Supprime les préfixes textuels (A., (B), -, •…)
        clean = re.sub(r"^([□☐○◦◻▪●•\-–]|[A-Ea-e][.)]|\([A-Ea-e]\))\s*", "", clean)
        return clean.strip()

    for line in lines:
        kind = _classify_line(line)

        if kind in ("NOISE", "SKIP"):
            continue

        # Barème récapitulatif → arrêter le parsing des questions
        if kind == "SUMMARY":
            _save()
            current_q = None
            # Ignorer tout le reste (barème ne fait pas partie du formulaire)
            break

        if kind == "SECTION":
            _save()
            current_section    = line
            current_subsection = ""
            in_vraifaux_section = False
            current_context    = ""
            collecting_context = False
            collecting_travail = False
            travail_table_started = False
            continue

        if kind == "SUBSECTION":
            # ── "Travail à Faire" : étude de cas basée sur un tableau à compléter ──
            # (ex : Section "Maîtrise des risques des SMQ") — il n'y a pas de questions
            # numérotées : on capture les consignes comme une question ouverte unique,
            # et le tableau/légende est rendu séparément en HTML fidèle (table_html)
            # afin d'être affiché exactement comme dans le document source.
            if re.search(r"Travail\s+.{1,2}\s*[Ff]aire", line, re.IGNORECASE):
                _save()
                current_subsection = line
                current_q          = _new_q(line, "open")
                if case_study_table_html:
                    current_q["table_html"] = case_study_table_html
                current_part       = None
                collecting_travail = True
                travail_table_started = False
                current_context    = ""
                collecting_context = False
                continue

            _save()
            current_subsection  = line
            # Détecter section Vrai/Faux
            in_vraifaux_section = bool(re.search(
                r"Vrai\s*/\s*Faux|Vraifaux|True\s*/\s*False",
                line, re.IGNORECASE
            ))
            current_context    = ""
            collecting_context = False
            collecting_travail = False
            travail_table_started = False
            continue

        # ── Bloc "Travail à Faire" : capturer les consignes telles quelles, mais
        # arrêter dès qu'on atteint la légende/tableau (reconstruits proprement en
        # HTML via table_html — le texte linéarisé brut serait, lui, garbled) ──
        if collecting_travail and current_q is not None:
            if not travail_table_started and re.search(
                r"Probabilit[ée2]?\s*\(\s*P\s*\)|^\s*P\s*:\s*Probabilit", line, re.IGNORECASE
            ):
                travail_table_started = True
            if travail_table_started:
                continue
            current_q["text"] += "\n" + line
            continue

        # ── "Contexte N : …" — étude de cas : texte affiché avant les questions liées ──
        if re.match(r"^Contexte\s*\d+\s*[:.]", line, re.IGNORECASE):
            _save()
            current_context    = line
            collecting_context = True
            continue

        if kind == "QUESTION":
            _save()
            clean_text = re.sub(r"^(\d{1,2})[.)]\s*", "", line).strip()
            if current_context:
                clean_text = f"{current_context}\n\n{clean_text}"
            collecting_context = False
            if in_vraifaux_section:
                # Question composée : main text + sous-questions a/b/c à venir
                current_q   = _new_q(clean_text, "compound")
                current_part = None
            else:
                in_qcm = bool(re.search(r"QCM|Choix\s+multiple|MCQ",
                                         current_subsection, re.IGNORECASE))
                current_q   = _new_q(clean_text, "qcm" if in_qcm else "open")
                current_part = None
            continue

        # ── Ligne "Vrai ☐ Faux ☐" → marquer la sous-question courante comme vraifaux ──
        if re.match(r"^Vrai\s*[☐□]", line, re.IGNORECASE) and current_q is not None:
            if current_q.get("type") == "compound" and current_part:
                current_part["type"] = "vraifaux"
            continue

        # ── Sous-questions a. / b. / c. dans une question composée ──
        sub_match = re.match(r"^([a-e])[.)]\s+(.+)", line) if current_q else None
        if sub_match and current_q and current_q.get("type") == "compound":
            # Sauvegarder la sous-question précédente
            if current_part:
                current_q["parts"].append(current_part)
            label_letter = sub_match.group(1)
            label_text   = sub_match.group(2).strip()
            current_part = {
                "id":    str(uuid.uuid4()),
                "label": f"{label_letter}. {label_text}",
                "type":  "open",
            }
            continue

        # ── "Justifier votre réponse" dans une sous-question composée ──
        if kind == "JUSTIF" and current_q is not None:
            if current_q.get("type") == "compound" and current_part:
                # C'est le label d'une sous-question "Justifier"
                if current_part:
                    current_q["parts"].append(current_part)
                current_part = {
                    "id":    str(uuid.uuid4()),
                    "label": line,
                    "type":  "open",
                }
            else:
                current_q["has_justification"] = True
            continue

        if kind == "OPTION" and current_q is not None:
            if current_q.get("type") != "compound":
                opt = _opt_text(line)
                if opt:
                    current_q["options"].append(opt)
            continue

        # ── Accumulation du texte de contexte (paragraphe sous "Contexte N :") ──
        if collecting_context:
            # Exclure l'en-tête "Questions d'évaluation" qui suit le paragraphe
            if not re.match(r"^Questions?\s+d.[ée]valuation", line, re.IGNORECASE):
                current_context += "\n" + line
            continue

        # TEXT : continuation du texte de question
        if current_q is not None:
            if re.match(r"^\(?\d+[\.,]?\d*\s*points?\)?\.?$", line, re.IGNORECASE):
                continue
            # Dans une question composée, le texte continue la question principale
            if current_q.get("type") == "compound" and not current_part:
                current_q["text"] += "\n" + line
            elif current_q.get("type") != "compound":
                current_q["text"] += "\n" + line

    # Sauvegarder la dernière sous-question si composée
    if current_q and current_q.get("type") == "compound" and current_part:
        current_q["parts"].append(current_part)
    _save()
    return questions
